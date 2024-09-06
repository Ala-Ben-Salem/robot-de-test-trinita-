import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException


def wait_for_element(driver, by, value, timeout=15):
    """Wait for an element to be present and visible."""
    try:
        wait = WebDriverWait(driver, timeout)
        element = wait.until(EC.visibility_of_element_located((by, value)))
        return element
    except TimeoutException:
        print(f"Timeout: L'element avec {by}='{value}' n'a pas ete trouve apres {timeout} secondes.")
        return None

def capture_screenshot(driver, step_name, screenshot_dir='screenshots'):
    """Capture a screenshot and save it to the screenshots directory."""
    if not os.path.exists(screenshot_dir):
        os.makedirs(screenshot_dir)
    screenshot_path = os.path.join(screenshot_dir, f"{step_name}.png")
    driver.save_screenshot(screenshot_path)
    print(f"Screenshot saved: {screenshot_path}")


def add_rapport_word(rapport_folder, text, image_path, compteur):
    # Vérifier si le dossier 'rapport' existe, sinon le créer
    if not os.path.exists(rapport_folder):
        os.makedirs(rapport_folder)

    # Chemin complet pour le fichier Word
    rapport_path = os.path.join(rapport_folder, f'final_report_{compteur}.docx')

    #
    if os.path.exists(rapport_path):
        # Ouvrir le document Word existant
        doc = Document(rapport_path)
    else:
        # Créer un nouveau document Word
        doc = Document()

        #
        title = doc.add_heading('PROJET STAGE D\'IMMERSION', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #
        subtitle1 = doc.add_heading('DIPLÔME NATIONAL D\'INGÉNIEUR', level=2)
        subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle2 = doc.add_heading('SPÉCIALITÉ : INFORMATIQUE', level=3)
        subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #
        doc.add_paragraph('2024 - 2025', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Robot de test Trinita selenium', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER

        #
        doc.add_paragraph('Réalisé par : Ben Salem Ala', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Encadrant Entreprise : Rihem Wadaa', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_heading('Rapport Final des Captures d\'Écran', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # Ajouter le texte
    doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter l'image
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6))
    else:
        doc.add_paragraph('Image non trouvée ou n\'a pas été prise.')

    # Ajouter un espace après chaque ajout
    doc.add_paragraph()

    # Sauvegarder le document Word
    doc.save(rapport_path)
    print(f"Le rapport a été sauvegardé sous {rapport_path}")



def clear_screenshots_folder(folder_path='screenshots'):
    """Vide le dossier spécifié en supprimant tous les fichiers qu'il contient."""
    # Vérifie si le dossier existe
    if os.path.exists(folder_path):
        # Parcourt tous les fichiers dans le dossier
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            # Vérifie si c'est un fichier
            if os.path.isfile(file_path):
                os.remove(file_path)  # Supprime le fichier
        print(f"Le dossier '{folder_path}' a été vidé.")
    else:
        print(f"Le dossier '{folder_path}' n'existe pas.")

